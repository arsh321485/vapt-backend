"""
File parser utilities for scope targets.
Supports: Excel (.xlsx, .xls), CSV (.csv), Word (.docx), Text (.txt)
"""
import os
import re
from typing import List, Dict, Optional
import pandas as pd

# Try to import python-docx for Word files
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

from .utils import extract_targets_from_text, classify_target, is_valid_subnet, expand_subnet


def parse_excel(file_path: str, expand_subnets: bool = True) -> List[Dict[str, any]]:
    """
    Parse Excel file (.xlsx, .xls) and extract IPs/URLs/Subnets from all columns.
    
    Args:
        file_path: Path to Excel file
        expand_subnets: If True, expand subnets into individual IPs. If False, keep as subnet with count.
    
    Returns:
        List of dicts with 'target_type', 'target_value', and optionally 'subnet_count'
    """
    targets = []
    seen = set()
    
    try:
        # Read Excel file
        df = pd.read_excel(file_path, header=None)
        
        # Collect all values into text format (one per line)
        text_lines = []
        for column in df.columns:
            for value in df[column].dropna():
                value_str = str(value).strip()
                if value_str and value_str.lower() not in ['nan', 'none', '']:
                    text_lines.append(value_str)
        
        # Use extract_targets_from_text for consistent processing
        if text_lines:
            text_content = '\n'.join(text_lines)
            extracted = extract_targets_from_text(text_content, expand_subnets=expand_subnets)
            # Deduplicate
            for item in extracted:
                if item['target_value'] not in seen:
                    seen.add(item['target_value'])
                    targets.append(item)
    
    except Exception as e:
        raise ValueError(f"Error parsing Excel file: {str(e)}")
    
    return targets


def parse_csv(file_path: str, expand_subnets: bool = True) -> List[Dict[str, any]]:
    """
    Parse CSV file and extract IPs/URLs/Subnets from all columns.
    
    Args:
        file_path: Path to CSV file
        expand_subnets: If True, expand subnets into individual IPs. If False, keep as subnet with count.
    
    Returns:
        List of dicts with 'target_type', 'target_value', and optionally 'subnet_count'
    """
    targets = []
    seen = set()
    
    try:
        # Read CSV file (try different encodings)
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, header=None, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if df is None:
            raise ValueError("Could not decode CSV file with any supported encoding")
        
        # Collect all values into text format (one per line)
        text_lines = []
        for column in df.columns:
            for value in df[column].dropna():
                value_str = str(value).strip()
                if value_str and value_str.lower() not in ['nan', 'none', '']:
                    text_lines.append(value_str)
        
        # Use extract_targets_from_text for consistent processing
        if text_lines:
            text_content = '\n'.join(text_lines)
            extracted = extract_targets_from_text(text_content, expand_subnets=expand_subnets)
            # Deduplicate
            for item in extracted:
                if item['target_value'] not in seen:
                    seen.add(item['target_value'])
                    targets.append(item)
    
    except Exception as e:
        raise ValueError(f"Error parsing CSV file: {str(e)}")
    
    return targets


def parse_word(file_path: str) -> List[Dict[str, str]]:
    """
    Parse Word document (.docx) and extract IPs/URLs from text.
    
    Args:
        file_path: Path to Word file
    
    Returns:
        List of dicts with 'target_type' and 'target_value'
    """
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx library is not installed. Install it with: pip install python-docx")
    
    targets = []
    seen = set()
    
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)
        
        # Combine all text and extract targets
        full_text = '\n'.join(text_content)
        extracted = extract_targets_from_text(full_text)
        
        # Deduplicate
        for item in extracted:
            if item['target_value'] not in seen:
                seen.add(item['target_value'])
                targets.append(item)
    
    except Exception as e:
        raise ValueError(f"Error parsing Word file: {str(e)}")
    
    return targets


def parse_text(file_path: str, expand_subnets: bool = True) -> List[Dict[str, any]]:
    """
    Parse text file (.txt) and extract IPs/URLs/Subnets.
    
    Args:
        file_path: Path to text file
        expand_subnets: If True, expand subnets into individual IPs. If False, keep as subnet with count.
    
    Returns:
        List of dicts with 'target_type', 'target_value', and optionally 'subnet_count'
    """
    targets = []
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError("Could not decode text file with any supported encoding")
        
        # Extract targets from text
        targets = extract_targets_from_text(content, expand_subnets=expand_subnets)
    
    except Exception as e:
        raise ValueError(f"Error parsing text file: {str(e)}")
    
    return targets


def parse_file(file_path: str, expand_subnets: bool = True) -> List[Dict[str, any]]:
    """
    Main dispatcher function that routes to appropriate parser based on file extension.
    
    Args:
        file_path: Full path to the file
        expand_subnets: If True, expand subnets into individual IPs. If False, keep as subnet with count.
    
    Returns:
        List of dicts with 'target_type', 'target_value', and optionally 'subnet_count'
    
    Raises:
        ValueError: If file type is unsupported or parsing fails
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    # Excel files
    if ext in ('.xlsx', '.xls'):
        return parse_excel(file_path, expand_subnets=expand_subnets)
    
    # CSV files
    if ext == '.csv':
        return parse_csv(file_path, expand_subnets=expand_subnets)
    
    # Word files
    if ext in ('.docx', '.doc'):
        return parse_word(file_path)
    
    # Text files
    if ext == '.txt':
        return parse_text(file_path, expand_subnets=expand_subnets)
    
    raise ValueError(f"Unsupported file type: {ext}. Supported: .xlsx, .xls, .csv, .docx, .doc, .txt")
